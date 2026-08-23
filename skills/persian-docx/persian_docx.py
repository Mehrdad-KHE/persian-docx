"""Persian-correct .docx builder. python-docx has no RTL support, so the bidi
flags, the complex-script font and the logical alignment are written into the
XML here. Nothing in this file is machine-specific.

    from persian_docx import PersianDoc
    d = PersianDoc()
    d.title("قرارداد"); d.h1("بخش اول"); d.p("متن ..."); d.bullet("یک")
    d.save("out.docx")

Fonts: Vazirmatn (SIL OFL, free) is the default for both body and headings.
Set TITLE_FONT, or PERSIAN_DOCX_TITLE_FONT in the environment, to use a display
face you have a licence for. fonts_present() reports what is actually installed.
"""
import os
import re
import glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = os.environ.get("PERSIAN_DOCX_FONT", "Vazirmatn")
TITLE_FONT = os.environ.get("PERSIAN_DOCX_TITLE_FONT", FONT)
FALLBACK = "Tahoma"
BULLET = "•"          # Word's own list renders as "/" in RTL Persian, and
                            # an arrow glyph reads as Word's collapse marker

LATIN_DEFAULTS = ("Calibri", "Times New Roman", "Arial", "Cambria",
                  "Aptos", "Segoe UI", "")

VAZIRMATN_URL = "https://github.com/rastikerdar/vazirmatn/releases"


def fonts_present(*names):
    """Which of these font families are installed on this machine (Windows,
    macOS, Linux). A missing font is not an error — Word substitutes — but the
    result stops looking Persian, so callers should warn."""
    names = names or (FONT, TITLE_FONT)
    roots = [
        os.path.expandvars(r"%LOCALAPPDATA%\Microsoft\Windows\Fonts"),
        r"C:\Windows\Fonts",
        os.path.expanduser("~/Library/Fonts"), "/Library/Fonts",
        os.path.expanduser("~/.local/share/fonts"), "/usr/share/fonts",
    ]
    found = set()
    for root in roots:
        if not os.path.isdir(root):
            continue
        files = glob.glob(os.path.join(root, "**", "*.tt[fc]"), recursive=True)
        files += glob.glob(os.path.join(root, "**", "*.otf"), recursive=True)
        for f in files:
            base = os.path.basename(f).lower().replace(" ", "")
            for n in names:
                if n.lower().replace(" ", "").rstrip("s") in base:
                    found.add(n)
    return {n: (n in found) for n in names}


# Free to use and redistribute (SIL OFL 1.1), verified August 2026. Anything
# outside this list — IRANSans, IRANYekan, Dana, Yekan Bakh, Kalameh, Morabba,
# Ravi — is sold by fontiran.com and must be licensed. The "B" family (B Titr,
# B Nazanin, B Mitra, B Lotus) carries a copyright notice with no obtainable
# licence: it ships with pirated Office packs across Iran, so never embed it in
# a document that leaves the building.
FREE_FONTS = {
    "Vazirmatn": "body and headings; 9 weights; actively maintained",
    "Estedad":   "body; tighter than Vazirmatn, good for dense legal text",
    "Sahel":     "body; wider counters",
    "Lalezar":   "headings only; single weight, high contrast",
}

# Body size and leading by document type. Persian is set 1-3 points larger than
# Latin at the same apparent size - the letterforms sit smaller inside the em
# and the dots that tell letters apart are what you lose first. Leading is
# larger too: Persian ascenders and descenders reach further than Latin ones
# (W3C alreq), so 1.0 collides.
KINDS = {
    "contract": (12, 1.5),   # printed, signed, read by a notary or a lawyer
    "letter":   (13, 1.5),   # نامهٔ اداری
    "thesis":   (13, 1.5),   # Irandoc راه, table 3-1
    "book":     (13.5, 1.3),  # narrower measure, tighter leading
    "report":   (12, 1.5),
}


NOTE_COLOR = RGBColor(0xC0, 0x00, 0x00)


_EN_MONTHS = ("January", "February", "March", "April", "May", "June", "July",
              "August", "September", "October", "November", "December")
_FA_MONTHS = ("ژانویه", "فوریه",
              "مارس", "آوریل",
              "مه", "ژوئن", "ژوئیه",
              "آگوست", "سپتامبر",
              "اکتبر", "نوامبر",
              "دسامبر")
_FR_MONTHS = ("janvier", "février", "mars", "avril", "mai", "juin",
              "juillet", "août", "septembre", "octobre", "novembre",
              "décembre")


# Persian speakers write the English month names as they say them, and the
# spelling from a textbook is not the one that turns up in a real document.
_FA_ALIASES = {
    "ژانویه": 0, "جانویه": 0, "ژانویه‌": 0,
    "فوریه": 1, "فبروری": 1,
    "مارس": 2, "مارچ": 2,
    "آوریل": 3, "آپریل": 3,
    "مه": 4, "می": 4, "مای": 4,
    "ژوئن": 5, "جون": 5,
    "ژوئیه": 6, "جولای": 6, "جولایی": 6,
    "آگوست": 7, "اوت": 7, "آگست": 7,
    "سپتامبر": 8, "سپتمبر": 8,
    "اکتبر": 9, "اکتوبر": 9,
    "نوامبر": 10, "نومبر": 10,
    "دسامبر": 11, "دسمبر": 11,
}


_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")


def fa_quotes(text):
    """Straight quotes are wrong in Persian. Turn "..." into «...»."""
    out, opening = [], True
    for ch in text:
        if ch == '"':
            out.append("«" if opening else "»")
            opening = not opening
        else:
            out.append(ch)
    return "".join(out)


def fa_digits(text):
    """Persian digits — but not everywhere.

    Latin digits stay Latin in four places: inside a [...] placeholder, in an
    amount written with a currency sign, in an English date, and anywhere a
    Latin word owns the number (a version, a figure, a file name). A date
    reading "August ۲۱" is the worst of both languages and cannot be
    checked against its English original.
    """
    protect = [
        r"\[[^\]]*\]",
        r"[$€£]\s?[\d,.]+",
        r"(?:" + "|".join(_EN_MONTHS) + r")\s+\d{1,2}(?:\s*,)?\s*\d{4}",
        r"[A-Za-z][A-Za-z.\-]*\s*\d[\d,.:/\-]*|\d+\.\d[\d.]*",
        r"\d[\d,.:/\-]*\s*[A-Za-z]",
    ]
    spans = []
    for pat in protect:
        spans += [m.span() for m in re.finditer(pat, text)]

    def kept(i):
        return any(a <= i < b for a, b in spans)

    return "".join(ch if kept(i) else ch.translate(_DIGITS)
                   for i, ch in enumerate(text))


ZWNJ = "‌"

_LETTERS = {"ي": "ی", "ك": "ک", "ة": "ه",
            "٠": "۰", "١": "۱", "٢": "۲",
            "٣": "۳", "٤": "۴", "٥": "۵",
            "٦": "۶", "٧": "۷", "٨": "۸",
            "٩": "۹"}

_PREFIX = ("می", "نمی")
_SUFFIX = ("ها", "های", "هایی",
           "تر", "ترین")


def normalize(text):
    """Everything the machine can decide on its own: Arabic letters that look
    Persian, Arabic-Indic digits, the wrong comma and question mark, straight
    quotes, doubled spaces, a space before punctuation, and the half-space that
    Persian words need in order not to fuse."""
    for a, b in _LETTERS.items():
        text = text.replace(a, b)
    text = fa_quotes(text)
    text = text.replace(",", "،").replace(";", "؛").replace("?", "؟")
    # A comma standing between two Latin characters belongs to the Latin text
    # around it - a street address or an English date - and a Persian comma
    # there is simply wrong: "8185 Yonge Street، Unit 201".
    text = re.sub(r"(?<=[A-Za-z0-9])،(?=\s*[A-Za-z0-9])", ",", text)
    for pre in _PREFIX:
        text = re.sub(r"(?<![\w؀-ۿ])" + pre + r" (?=[ء-ی])",
                      pre + ZWNJ, text)
    for suf in _SUFFIX:
        text = re.sub(r"(?<=[؀-ۿ]) " + suf +
                      r"(?![\w؀-ۿ])", ZWNJ + suf, text)
    text = re.sub(r"\s+([،؛؟.:!])", r"\1", text)
    text = re.sub(r"\.{3,}", "…", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    return text


def check_text(text):
    """What normalize() cannot decide — reported, not changed."""
    out = []
    if re.search(r"می‌?باشد", text):
        out.append("«می‌باشد» — use «است»")
    for word in ("بدینوسیله",
                 "در راستای",
                 "مورد بررسی"):
        if word in text:
            out.append("translated phrasing: " + word)
    if text.count("—") > 1:
        out.append("more than one em dash in one passage")
    if re.search(r"\d,\d{3}", text) and "$" not in text:
        out.append("Latin thousands separator in Persian prose — use ٬")
    return out



# ---------------------------------------------------------------- audience --

IRAN_MARKERS = (
    "تهران", "ایران",           # تهران، ایران
    "ریال", "تومان",                   # ریال، تومان
    "قانون مدنی",                      # قانون مدنی
    "کد ملی", "شماره ملی",
    "ثبت احوال", "شمسی",
    "دفترخانه",
)

ABROAD_MARKERS = (
    "Ontario", "Canada", "CRA", "GST", "HST", "CAD", "USD", "EUR", "LLC",
    "Inc.", "GmbH", "کانادا", "تورنتو",
    "انتاریو", "دلار", "یورو",
)




def detect_audience(text):
    """Is this document going to be read inside Iran or outside it? The answer
    changes the month names, the currency and the calendar. Returns
    "iran", "abroad", or "unknown" — and unknown means ask, never guess."""
    lower = text.lower()
    iran = sum(1 for m in IRAN_MARKERS if m in text)
    abroad = sum(1 for m in ABROAD_MARKERS if m.lower() in lower)
    if iran > abroad and iran:
        return "iran"
    if abroad > iran and abroad:
        return "abroad"
    return "unknown"


def month_style(audience):
    """Which month names belong in this document.

    iran    → Persian names, and a Jalali date is what the reader expects.
    abroad  → the month name in the language of the country the paper lives in.
              A Persian month on a document a foreign office must act on cannot
              be checked against its source, which is worse than useless.
    Never French unless the document is genuinely for Quebec or France; Canada
    being bilingual is not a reason.
    """
    return {"iran": "fa", "abroad": "en"}.get(audience, "ask")


def convert_months(text, to="en"):
    """Translate month names between Persian, English and French. Only the
    names — a date's digits and order are left alone, because a date that stops
    matching its source is a date nobody can verify."""
    tables = {"fa": _FA_MONTHS, "en": _EN_MONTHS, "fr": _FR_MONTHS}
    target = tables[to]
    # A month name must stand as its own word. Without this, Persian "مه" (May)
    # eats the middle of "مهلت" and the document quietly turns to nonsense.
    letters = r"A-Za-zÀ-ſ؀-ۿ"
    for group in tables.values():
        if group is target:
            continue
        for i, name in enumerate(group):
            pat = r"(?<![" + letters + r"])" + re.escape(name) + r"(?![" + letters + r"])"
            text = re.sub(pat, target[i], text, flags=re.IGNORECASE)
    return text



_TO_EN_DIGITS = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")


_FA_DIGITS_RE = "[۰-۹]"


def fa_date(text):
    """Persian word order for a date a Persian reader reads.

    In Persian the day comes first and the year is introduced: ۲۳ام آگوست
    سال ۲۰۲۶. Anything less than the whole date - month name, order,
    digits - leaves a hybrid that belongs to neither language.
    """
    names = "|".join(_EN_MONTHS + _FA_MONTHS + tuple(_FA_ALIASES))
    d = r"[0-9۰-۹]"

    def idx(mon):
        low = mon.lower()
        if low in _FA_ALIASES:
            return _FA_ALIASES[low]
        for group in (_EN_MONTHS, _FA_MONTHS):
            for i, n in enumerate(group):
                if n.lower() == low:
                    return i
        return None

    def build(mon, day, year):
        i = idx(mon)
        if i is None:
            return mon
        year = year.translate(_DIGITS)
        if day is None:
            return "%s سال %s" % (_FA_MONTHS[i], year)
        day = str(int(day.translate(_TO_EN_DIGITS))).translate(_DIGITS)
        return "%sام %s سال %s" % (day, _FA_MONTHS[i], year)

    # English order first: August 23, 2026
    pat_en = (r"(?P<mon>" + names + r")\s+(?P<day>" + d +
              r"{1,2})\s*[،,]?\s*(?P<year>" + d + r"{4})")
    text = re.sub(pat_en, lambda m: build(m.group("mon"), m.group("day"),
                                          m.group("year")), text, flags=re.I)
    # Persian order, possibly already half-converted: 23 August 2026
    pat_fa = (r"(?P<day>" + d + r"{1,2})\s*(?:ام)?\s+(?P<mon>" + names +
              r")\s+(?:سال\s+)?(?P<year>" + d + r"{4})")
    text = re.sub(pat_fa, lambda m: build(m.group("mon"), m.group("day"),
                                          m.group("year")), text, flags=re.I)
    # A month with no day still needs its own language: July 2026
    pat_my = (r"(?P<mon>" + names + r")\s+(?:سال\s+)?(?P<year>" + d + r"{4})")
    text = re.sub(pat_my, lambda m: build(m.group("mon"), None,
                                          m.group("year")), text, flags=re.I)
    return text


def en_date(text):
    """English form for a date on a document a foreign office will act on.

    Converting only the month name leaves "August ۲۳، ۲۰۲۶" - an English month,
    Persian digits and a Persian comma. A bank or a lawyer cannot check that
    against the English original.
    """
    names = "|".join(_EN_MONTHS + _FA_MONTHS + tuple(_FA_ALIASES))
    d = r"[0-9۰-۹]"

    def idx(mon):
        low = mon.lower()
        if low in _FA_ALIASES:
            return _FA_ALIASES[low]
        for group in (_EN_MONTHS, _FA_MONTHS):
            for i, n in enumerate(group):
                if n.lower() == low:
                    return i
        return None

    def build(mon, day, year):
        i = idx(mon)
        if i is None:
            return mon
        year = year.translate(_TO_EN_DIGITS)
        if day is None:
            return "%s %s" % (_EN_MONTHS[i], year)
        return "%s %d, %s" % (_EN_MONTHS[i], int(day.translate(_TO_EN_DIGITS)),
                              year)

    pat_fa = (r"(?P<day>" + d + r"{1,2})\s*(?:ام)?\s+(?P<mon>" + names +
              r")\s+(?:سال\s+)?(?P<year>" + d + r"{4})")
    text = re.sub(pat_fa, lambda m: build(m.group("mon"), m.group("day"),
                                          m.group("year")), text, flags=re.I)
    pat_en = (r"(?P<mon>" + names + r")\s+(?P<day>" + d +
              r"{1,2})\s*[،,]?\s*(?P<year>" + d + r"{4})")
    text = re.sub(pat_en, lambda m: build(m.group("mon"), m.group("day"),
                                          m.group("year")), text, flags=re.I)
    pat_my = (r"(?P<mon>" + names + r")\s+(?:سال\s+)?(?P<year>" + d + r"{4})")
    text = re.sub(pat_my, lambda m: build(m.group("mon"), None,
                                          m.group("year")), text, flags=re.I)
    return text


def mixed_dates(text):
    """Report dates that mix the two languages — an English month name sitting
    beside Persian digits, or a Persian month beside Latin ones."""
    out = []
    names = "|".join(_EN_MONTHS)
    for m in re.finditer(r"(?:" + names + r")\s*" + _FA_DIGITS_RE, text, re.I):
        out.append("English month with Persian digits: " + m.group(0))
    for m in re.finditer(r"(?:" + "|".join(_FA_MONTHS) + r")\s*[0-9]", text):
        out.append("Persian month with Latin digits: " + m.group(0))
    return out


def _rtl_paragraph(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _jc(p, val):
    """Set alignment by hand. In a bidi paragraph Word reads "right" as the
    logical right, which lands on the physical LEFT — headings written that way
    hug the wrong margin. "start" is the one that means the reader's side."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)


def _rtl_run(run, font=None):
    font = font or FONT
    rPr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:cs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "1")
        rPr.append(el)
    # complex-script font must be set separately or Word ignores it for Persian
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


class PersianDoc:
    def __init__(self, title=None, digits=True, audience=None, kind="report",
                 dates=None):
        """audience: "iran", "abroad", or None to leave month handling alone.
        It decides month names and nothing else — see month_style().
        kind: which document this is — see KINDS. It sets body size and
        leading, the two numbers Persian readers notice first.
        dates: "fa" or "en", overriding audience. The date follows the language
        the document is READ in, not the country it is filed in. A Persian
        translation of an Ontario contract is read in Persian — its dates
        belong in Persian, even though the English original that gets signed
        keeps August 23, 2026."""
        self.doc = Document()
        self.digits = digits
        self.audience = audience
        self.dates = dates
        self.size, self.leading = KINDS.get(kind, KINDS["report"])
        st = self.doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(self.size)
        st.element.rPr.rFonts.set(qn("w:cs"), FONT)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK)
        for s in self.doc.sections:
            s.right_margin = s.left_margin = Cm(2.5)
            sectPr = s._sectPr
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)
        if title:
            self.h1(title)

    def _style_rtl(self, name, size=None, bold=None, space_before=None,
                   font=None):
        """Word's built-in styles are LTR and use their own Latin font. A heading
        that is not fixed here comes out left-aligned in the wrong typeface."""
        font = font or FONT
        st = self.doc.styles[name]
        st.font.name = font
        st.element.rPr.rFonts.set(qn("w:cs"), font)
        st.element.rPr.rFonts.set(qn("w:ascii"), font)
        st.element.rPr.rFonts.set(qn("w:hAnsi"), font)
        if size:
            st.font.size = Pt(size)
        if bold is not None:
            st.font.bold = bold
        st.font.color.rgb = None
        pf = st.paragraph_format
        if space_before:
            pf.space_before = Pt(space_before)
        pPr = st.element.get_or_add_pPr()
        if pPr.find(qn("w:bidi")) is None:
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.append(bidi)
        return st

    def _add(self, text, style=None, size=None, bold=False, space_after=8,
             align="start", font=None):
        p = self.doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = self.leading
        _rtl_paragraph(p)
        _jc(p, align)
        text = normalize(text)
        style = self.dates or (month_style(self.audience)
                               if self.audience in ("iran", "abroad") else None)
        if style in ("fa", "en"):
            text = convert_months(text, to=style)
            # Half a conversion is worse than none: the month name alone leaves
            # a date in two languages that matches neither source.
            text = fa_date(text) if style == "fa" else en_date(text)
        run = p.add_run(fa_digits(text) if self.digits else text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        _rtl_run(run, font)
        return p

    def h1(self, text):
        self._style_rtl("Heading 1", size=18, bold=False, space_before=18,
                        font=TITLE_FONT)
        return self._add(text, style="Heading 1", space_after=12,
                         font=TITLE_FONT)

    def h2(self, text):
        self._style_rtl("Heading 2", size=14, bold=False, space_before=14,
                        font=TITLE_FONT)
        return self._add(text, style="Heading 2", space_after=8,
                         font=TITLE_FONT)

    def h3(self, text):
        self._style_rtl("Heading 3", size=12, bold=True, space_before=10)
        return self._add(text, style="Heading 3", space_after=6)

    def title(self, text):
        self._style_rtl("Title", size=26, bold=False, font=TITLE_FONT)
        return self._add(text, style="Title", space_after=18, font=TITLE_FONT,
                         align="center")

    def p(self, text, justify=True):
        """Body text is justified — a ragged right edge is what makes a Persian
        page look unfinished. Pass justify=False for a short standalone line."""
        return self._add(text, align="both" if justify else "start")

    def bullet(self, text, justify=True):
        """Word's own bullet list comes out as "/" in an RTL Persian document,
        so the marker is written into the text and the indent set by hand."""
        p = self._add(BULLET + " " + text,
                      align="both" if justify else "start", space_after=4)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        return p

    def note(self, text):
        """A red explanatory line under a clause: what this clause actually does
        to the reader, in plain Persian. It is commentary, not contract text —
        the colour and the indent are what tell the two apart on paper."""
        p = self._add("◄ " + text, align="both", space_after=10,
                      size=self.size - 1.5)
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.right_indent = Cm(1.2)
        for r in p.runs:
            r.font.color.rgb = NOTE_COLOR
            r.italic = True
        return p

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        return path


def demo(path):
    d = PersianDoc()
    d.title("سندِ آزمایشی")
    d.h1("بخشِ اول — عنوانِ سطحِ یک")
    d.h2("۱. عنوانِ سطحِ دو")
    d.p("این یک جملهٔ فارسی با اصطلاحِ انگلیسیِ escrow و عددِ 1450 است.")
    d.h3("۱.۱ عنوانِ سطحِ سه")
    d.bullet("بندِ اول")
    d.bullet("بندِ دوم با جای خالی [مبلغ] و تاریخِ August 21, 2026")
    return d.save(path)


def audit(path):
    """Open a finished .docx and report what is wrong for a Persian reader."""
    doc = Document(path)
    issues, heads = [], 0
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        if p.style.name.startswith(("Heading", "Title")):
            heads += 1
        pPr = p._p.pPr
        if pPr is None or pPr.find(qn("w:bidi")) is None:
            issues.append((i, "no w:bidi", p.text[:40]))
        pPr = p._p.pPr
        jc = None if pPr is None else pPr.find(qn("w:jc"))
        val = None if jc is None else jc.get(qn("w:val"))
        if val not in ("start", "both", "center"):
            issues.append((i, "alignment %s — use start/both/center" % val,
                           p.text[:40]))
        for bad in mixed_dates(p.text):
            issues.append((i, bad, p.text[:40]))
        if '"' in p.text:
            issues.append((i, 'straight quote — use « »', p.text[:40]))
        for r in p.runs[:1]:
            rPr = r._element.rPr
            if rPr is None or rPr.find(qn("w:rtl")) is None:
                issues.append((i, "run not rtl", p.text[:40]))
            else:
                rf = rPr.find(qn("w:rFonts"))
                cs = None if rf is None else rf.get(qn("w:cs"))
                # Any face is fine if it was chosen deliberately. A Latin
                # default here means the complex-script slot was never set and
                # every Persian glyph silently falls back to Word's own font.
                if not cs or cs in LATIN_DEFAULTS:
                    issues.append((i, "no complex-script font (cs=%s)" % cs,
                                   p.text[:40]))
    if heads == 0:
        issues.append((0, "no real headings — Word sees no structure", ""))
    return issues


if __name__ == "__main__":
    import sys
    print(demo(sys.argv[1] if len(sys.argv) > 1 else "persian-docx-demo.docx"))
